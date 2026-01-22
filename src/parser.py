import os
import re
from typing import List, Dict, Union, Any
import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import mammoth
import io

class DocxParser:
    def __init__(self, temp_image_dir: str = "temp_images"):
        self.temp_image_dir = temp_image_dir
        if not os.path.exists(temp_image_dir):
            os.makedirs(temp_image_dir)

    def get_content_and_html(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX and return both structured content and tagged HTML.
        Uses a robust sibling-tagging strategy to avoid breaking Word fields.
        """
        from bs4 import BeautifulSoup
        
        # 1. Extract clean content for LLM using a fresh doc object
        doc = docx.Document(file_path)
        content_list = []
        element_id = 0
        
        for element in doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, doc)
                images = self._extract_images_from_paragraph(paragraph, doc)
                for img_path in images:
                    content_list.append({"id": element_id, "type": "image", "path": img_path})
                    element_id += 1
                
                # 使用更可靠的方法提取段落文本
                # paragraph.text在某些情况下会失败，改用直接从XML提取
                text = self._extract_paragraph_text(element).strip()
                if text:
                    content_list.append({"id": element_id, "type": "text", "content": text})
                    element_id += 1
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_data = self._extract_table_data(table)
                content_list.append({"id": element_id, "type": "table", "content": table_data})
                element_id += 1
            else:
                # 处理SDT (Structured Document Tag) 元素，如目录等
                tag = element.tag
                if 'sdt' in tag.lower():
                    sdt_texts = self._extract_sdt_content(element, doc)
                    for text in sdt_texts:
                        if text.strip():
                            content_list.append({"id": element_id, "type": "text", "content": text})
                            element_id += 1

        # 2. Create a marked version for HTML conversion
        # We insert marker paragraphs BEFORE each target element
        marked_doc = docx.Document(file_path)
        element_id = 0
        body_elements = list(marked_doc.element.body)
        for element in body_elements:
            target_id = None
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, marked_doc)
                images = self._extract_images_from_paragraph(paragraph, marked_doc)
                element_id += len(images)
                # 使用更可靠的文本提取方法
                text = self._extract_paragraph_text(element).strip()
                if text:
                    target_id = element_id
                    element_id += 1
            elif isinstance(element, CT_Tbl):
                target_id = element_id
                element_id += 1
            else:
                # 处理SDT元素
                tag = element.tag
                if 'sdt' in tag.lower():
                    sdt_texts = self._extract_sdt_content(element, marked_doc)
                    # 为SDT中的每个文本项创建标记
                    for text in sdt_texts:
                        if text.strip():
                            target_id = element_id
                            # 为SDT中的每个段落添加标记
                            new_p = marked_doc.add_paragraph(f"MARKER_ID_{target_id}")
                            element.addprevious(new_p._p)
                            element_id += 1
                    continue  # 跳过后续的标记插入
            
            if target_id is not None:
                # Insert a marker paragraph before this element
                new_p = marked_doc.add_paragraph(f"MARKER_ID_{target_id}")
                element.addprevious(new_p._p)

        # 3. Convert to HTML
        doc_buffer = io.BytesIO()
        marked_doc.save(doc_buffer)
        doc_buffer.seek(0)
        html_res = mammoth.convert_to_html(doc_buffer)
        soup = BeautifulSoup(html_res.value, 'html.parser')

        # 4. Post-process HTML: find markers and tag siblings
        for marker_p in soup.find_all('p', string=re.compile(r'MARKER_ID_\d+')):
            m = re.search(r'MARKER_ID_(\d+)', marker_p.get_text())
            if m:
                eid = m.group(1)
                # The actual element is the next sibling
                sibling = marker_p.find_next_sibling()
                if sibling:
                    sibling['id'] = f"doc-el-{eid}"
                marker_p.decompose() # Remove the marker

        # 5. Tag images: match content_list images with HTML <img> tags
        img_tags = soup.find_all('img')
        img_content_items = [item for item in content_list if item['type'] == 'image']
        
        # Simple matching: assume order is preserved
        for img_tag, img_item in zip(img_tags, img_content_items):
            img_id = img_item['id']
            # Wrap img in a div for easier styling
            wrapper = soup.new_tag('div', id=f"doc-el-{img_id}", style="margin: 10px 0;")
            img_tag.wrap(wrapper)

        return {
            "content": content_list,
            "html": str(soup)
        }

    def _extract_images_from_paragraph(self, paragraph: Paragraph, doc: Document) -> List[str]:
        images = []
        for run in paragraph.runs:
            if 'drawing' in run.element.xml:
                xml = run.element.xml
                embeds = re.findall(r'r:embed="([^"]+)"', xml)
                for embed_id in embeds:
                    if embed_id in doc.part.rels:
                        rel = doc.part.rels[embed_id]
                        if "image" in rel.target_ref:
                            image_part = rel.target_part
                            image_data = image_part.blob
                            filename = f"img_{len(os.listdir(self.temp_image_dir))}.png"
                            filepath = os.path.join(self.temp_image_dir, filename)
                            with open(filepath, "wb") as f:
                                f.write(image_data)
                            images.append(filepath)
        return images

    def _extract_table_data(self, table: Table) -> str:
        """
        提取表格数据，包含标题和所有行（包括表头）
        Args:
            table: docx表格对象
            prev_text: 前一个段落文本，可能是表格标题
        Returns:
            格式化的表格字符串
        """
        result = []
        
        # 提取所有表格行（包括表头）
        if not table.rows:
            return "【空表格】"
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # 使用 itertext() 提取所有文本内容，包括嵌套元素
                # 这样可以获取到所有文本节点，即使 cell.text 为空
                # 去除重复的文本片段，保持顺序
                seen = []
                for text in cell._element.itertext():
                    text = text.strip()
                    if text and text not in seen:
                        seen.append(text)
                cell_text = '|'.join(seen).replace('\n', ' ')
                
                cells.append(cell_text)
            row_str = "| " + " | ".join(cells) + " |"
            result.append(row_str)
        
        return "\n".join(result)
    
    def _extract_paragraph_text(self, paragraph_element) -> str:
        """
        从段落元素中提取文本内容
        使用直接XML解析方法，比paragraph.text更可靠
        Args:
            paragraph_element: 段落XML元素
        Returns:
            段落文本
        """
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # 查找所有文本节点 (w:t)
        text_nodes = paragraph_element.findall('.//w:t', ns)
        
        # 合并所有文本
        text_parts = [node.text for node in text_nodes if node.text]
        
        return ''.join(text_parts)
    
    def _extract_sdt_content(self, sdt_element, doc: Document) -> List[str]:
        """
        提取SDT (Structured Document Tag)元素中的内容
        SDT通常用于目录、内容控件等
        Args:
            sdt_element: SDT元素
            doc: docx文档对象
        Returns:
            文本内容列表
        """
        texts = []
        
        # SDT的内容通常在 w:sdtContent 子元素中
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        sdt_content = sdt_element.find('.//w:sdtContent', ns)
        
        if sdt_content is not None:
            # 查找所有段落
            paragraphs = sdt_content.findall('.//w:p', ns)
            
            for p_element in paragraphs:
                # 提取段落中的所有文本节点
                text_nodes = p_element.findall('.//w:t', ns)
                p_text = ''.join([t.text for t in text_nodes if t.text]).strip()
                
                if p_text:
                    texts.append(p_text)
        
        return texts
